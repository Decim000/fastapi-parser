import io
import re
import os
import tempfile
import subprocess
import logging
import aspose.words as aw

from pathlib import Path
from docx import Document

from .zipsaver_service import ZipArchivedFileSaverService
from .s3_service import UploadToS3

class FileSaver:

    def __init__(self, title, bytes, number) -> None:
        self.title = title
        self.extension = self.title.split(".")[-1]
        self.file = bytes
        self.number = number

    def _save_file(self):
        logging.warning(f"{self.extension}")

        if (self.extension != "doc") and (self.extension != "zip") and (self.extension != "rar"):
            # logging.warning(f'i have a text from {self.title}')
            # data = io.BytesIO(self.file)
            # document = Document(data)
            # logging.warning(document.paragraphs[0].text)
            upload_service = UploadToS3(self.number, self.title, self.file)
            upload_service.upload_to_s3()
            

        if self.extension == "doc":
            # MAYBE use subprocess and libreoffice
            data = io.BytesIO(self.file)
            document = aw.Document(data)
            to_temp_path = Path.cwd().parent / "temp"
            docx_path = f"{os.path.dirname(to_temp_path)}/{self.title}x"
            document.save(docx_path, aw.SaveFormat.DOCX)
            # fp = tempfile.NamedTemporaryFile(suffix=".doc")
            # fp.write(self.file)
            # fp.seek(0)
            # temp_doc_path = fp.name
            # subprocess.call(['lowriter', '--headless', '--convert-to', 'docx', temp_doc_path])
            # fp.close()
        
        if self.extension == "zip":
            # if self.title.split(".")[-2] == "doc":
            with open(f"temp/{self.title}", "wb") as binary_file:
                binary_file.write(self.file)
                # binary_file.close()
            zip_saver = ZipArchivedFileSaverService(path = os.path.abspath(f"temp/{self.title}"), title = self.title, number = self.number)
            zip_saver._save_zipped_file()

